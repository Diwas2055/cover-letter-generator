import os
import subprocess
import shutil

from datetime import datetime
from docx import Document

root_path = os.path.abspath(os.curdir)
folder_path = root_path + "/templates/docx/"
output_dir = root_path + "/data/pdf/"


def main(
    user_name,
    job_title,
    company_name,
):
    if os.path.isdir(root_path + "/data/"):
        shutil.rmtree(root_path + "/data/")
    if os.path.isdir(output_dir):
        pass
    else:
        os.makedirs(output_dir)
    template_file_path = folder_path + f"{job_title.lower()}_cover_letter.docx"
    output_file_path = root_path + "/data/" + f"{user_name}-cover_letter.docx"

    current_date = datetime.now().strftime("%d %B, %Y")

    variables = {
        "${DATE}": current_date,
        "${JOB_TITLE}": job_title + "Developer",
        "${COMPANY_NAME}": company_name,
        "${YOUR_NAME}": user_name,
    }
    template_document = Document(template_file_path)
    try:
        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(
                                paragraph, variable_key, variable_value
                            )

        template_document.save(output_file_path)
        doc2pdf(output_file_path)
        if os.path.exists(output_file_path):
            os.remove(output_file_path)
            return {
                "file_location": output_dir + f"{user_name}-cover_letter.pdf",
                "file_name": f"{user_name}-cover_letter.pdf",
            }
    except Exception as e:
        raise e


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


try:
    from comtypes import client
except ImportError:
    client = None


def doc2pdf(doc):
    """
    convert a doc/docx document to pdf format
    :param doc: path to document
    """
    doc = os.path.abspath(doc)  # bugfix - searching files in windows/system32
    if client is None:
        return doc2pdf_linux(doc, output_dir)
    name, ext = os.path.splitext(doc)
    try:
        word = client.CreateObject("Word.Application")
        worddoc = word.Documents.Open(doc)
        worddoc.SaveAs(name + ".pdf", FileFormat=17)
    except Exception:
        raise
    finally:
        worddoc.Close()
        word.Quit()


def doc2pdf_linux(doc, output_path):
    """
    convert a doc/docx document to pdf format (linux only, requires libreoffice)
    :param doc: path to document
    """
    cmd = "libreoffice --convert-to pdf".split() + [doc]
    location_cmd = "--outdir".split() + [output_path]
    cmd.extend(location_cmd)
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    p.wait(timeout=10)
    stdout, stderr = p.communicate()
    if stderr:
        raise subprocess.SubprocessError(stderr)


def read_file(filename):
    file_path = root_path + "/templates/" + filename
    with open(file_path, "r") as file:
        # Read the contents of the file
        file_contents = file.read()
    return file_contents


def get_internship_template(data: dict):
    file_contents = read_file("intern_sample.txt")
    file_contents = file_contents.replace("${YOUR_NAME}", data["name"])
    file_contents = file_contents.replace("${COMPANY_NAME}", data["company"])
    file_contents = file_contents.replace("${STATUS}", data["status"])
    file_contents = file_contents.replace("${STUDY_FIELD}", data["study_field"])
    file_contents = file_contents.replace("${UNIVERSITY_NAME}", data["university_name"])

    return file_contents


# define Python user-defined exceptions
class InvalidChoiceException(Exception):
    "Raised when the input value is more than listed options"
    pass


if __name__ == "__main__":
    main("Diwash Bhandari", "Laravel", "Chuchuro Firm")
